#!/usr/bin/env python3
"""
最终修复知识库的脚本 - 解决所有已知问题
"""
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
import re
import shutil

# 尝试使用新的HuggingFace嵌入
try:
    from langchain_huggingface import HuggingFaceEmbeddings
    print("✅ 使用新版 langchain_huggingface.HuggingFaceEmbeddings")
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    print("⚠️ 回退到 langchain_community.embeddings.HuggingFaceEmbeddings")

# 配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'chroma_db'))
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

def extract_structured_content(docx_path):
    """提取结构化内容，保持上下文完整性"""
    print(f"📄 正在提取结构化内容...")
    
    doc = docx.Document(docx_path)
    
    # 分析文档结构
    content_blocks = []
    current_section = ""
    current_subsection = ""
    accumulated_content = ""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检测章节标题 (通常包含数字编号或特定关键词)
        if is_section_title(text):
            # 保存之前累积的内容
            if accumulated_content.strip():
                content_blocks.append({
                    'section': current_section,
                    'subsection': current_subsection,
                    'content': accumulated_content.strip()
                })
                accumulated_content = ""
            
            # 更新章节信息
            if is_major_section(text):
                current_section = text
                current_subsection = ""
            else:
                current_subsection = text
                
        else:
            # 累积正文内容
            if len(text) > 20:  # 过滤太短的文本
                accumulated_content += text + "\n"
    
    # 处理最后一块内容
    if accumulated_content.strip():
        content_blocks.append({
            'section': current_section,
            'subsection': current_subsection,
            'content': accumulated_content.strip()
        })
    
    print(f"✅ 提取了 {len(content_blocks)} 个内容块")
    return content_blocks

def is_section_title(text):
    """判断是否为章节标题"""
    # 常见的章节标题模式
    patterns = [
        r'^\d+[\.\s]',  # "1. " 或 "1 "
        r'^第\d+章',     # "第1章"
        r'^第\d+节',     # "第1节"
        r'^\d+\.\d+',   # "1.1"
        r'^[一二三四五六七八九十]+、',  # "一、"
    ]
    
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    
    # 如果文本很短且包含关键词，也可能是标题
    if len(text) < 50 and any(keyword in text for keyword in ['概述', '介绍', '方法', '技术', '原理', '应用']):
        return True
        
    return False

def is_major_section(text):
    """判断是否为主要章节"""
    return bool(re.match(r'^第\d+章|^\d+[\.\s]', text))

def create_contextual_documents(content_blocks):
    """创建包含上下文的文档"""
    print(f"📚 正在创建上下文丰富的文档...")
    
    documents = []
    
    for i, block in enumerate(content_blocks):
        content = block['content']
        section = block['section']
        subsection = block['subsection']
        
        # 为内容添加结构化的上下文
        contextual_content = ""
        
        if section:
            contextual_content += f"[章节: {section}]\n"
        if subsection:
            contextual_content += f"[小节: {subsection}]\n"
        
        contextual_content += content
        
        # 如果内容太长，进行进一步分割
        if len(contextual_content) > 1000:
            # 使用更智能的文本分割
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=150,
                separators=["\n\n", "。", "！", "？", "\n", "，", "、"],
                length_function=len,
            )
            
            # 创建临时文档进行分割
            temp_doc = Document(page_content=contextual_content, metadata={
                "source": DOCX_SOURCE_PATH,
                "section": section,
                "subsection": subsection,
                "block_id": i
            })
            
            sub_docs = text_splitter.split_documents([temp_doc])
            documents.extend(sub_docs)
        else:
            # 直接创建文档
            doc = Document(
                page_content=contextual_content,
                metadata={
                    "source": DOCX_SOURCE_PATH,
                    "section": section,
                    "subsection": subsection,
                    "block_id": i,
                    "length": len(contextual_content)
                }
            )
            documents.append(doc)
    
    print(f"✅ 创建了 {len(documents)} 个上下文文档")
    
    # 显示文档示例
    print(f"\n📋 文档示例:")
    for i, doc in enumerate(documents[:3]):
        print(f"\n文档 {i+1}:")
        print(f"  章节: {doc.metadata.get('section', '无')}")
        print(f"  小节: {doc.metadata.get('subsection', '无')}")
        print(f"  长度: {len(doc.page_content)} 字符")
        print(f"  内容: {doc.page_content[:150]}...")
    
    return documents

def rebuild_vector_database(documents):
    """重建向量数据库"""
    print(f"🚀 正在重建向量数据库...")
    
    # 删除旧数据库
    if os.path.exists(CHROMA_DB_PATH):
        print(f"🗑️ 删除旧的向量数据库...")
        shutil.rmtree(CHROMA_DB_PATH)
    
    # 加载嵌入模型
    print(f"🧠 加载嵌入模型: {EMBEDDING_MODEL_NAME}")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    
    # 创建新的向量数据库
    print(f"💾 创建新的向量数据库...")
    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=CHROMA_DB_PATH
    )
    
    print(f"✅ 向量数据库重建完成！")
    return vectordb

def comprehensive_test(vectordb):
    """全面测试新的知识库"""
    print(f"\n🧪 开始全面测试新知识库...")
    
    test_cases = [
        ("软件开发生命周期", "软件工程的核心流程"),
        ("需求分析方法", "软件需求获取和分析"),
        ("软件测试策略", "质量保证和测试方法"),
        ("面向对象设计", "OOP设计原理"),
        ("软件质量评价", "软件质量标准和度量"),
        ("项目管理技术", "软件项目的管理方法")
    ]
    
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})
    
    for query, description in test_cases:
        print(f"\n" + "="*60)
        print(f"🔍 测试查询: {query}")
        print(f"📝 期望内容: {description}")
        
        try:
            results = retriever.invoke(query)
            print(f"✅ 检索到 {len(results)} 个结果")
            
            for i, result in enumerate(results):
                print(f"\n📄 结果 {i+1}:")
                print(f"   章节: {result.metadata.get('section', '未知')}")
                print(f"   小节: {result.metadata.get('subsection', '未知')}")
                print(f"   内容长度: {len(result.page_content)} 字符")
                print(f"   内容预览: {result.page_content[:200]}...")
                
                # 计算相关性
                query_words = set(query.lower().split())
                content_words = set(result.page_content.lower().split())
                overlap = len(query_words.intersection(content_words))
                relevance = overlap / len(query_words) if query_words else 0
                print(f"   相关性得分: {relevance:.2f}")
                
        except Exception as e:
            print(f"❌ 测试失败: {e}")

def main():
    """主函数"""
    print("🚀 开始最终知识库修复...")
    
    try:
        # 检查文件是否存在
        if not os.path.exists(DOCX_SOURCE_PATH):
            print(f"❌ 找不到文件: {DOCX_SOURCE_PATH}")
            return
        
        # 1. 提取结构化内容
        content_blocks = extract_structured_content(DOCX_SOURCE_PATH)
        
        if not content_blocks:
            print("❌ 未能提取到有效内容")
            return
        
        # 2. 创建上下文文档
        documents = create_contextual_documents(content_blocks)
        
        if not documents:
            print("❌ 未能创建有效文档")
            return
        
        # 3. 重建向量数据库
        vectordb = rebuild_vector_database(documents)
        
        # 4. 全面测试
        comprehensive_test(vectordb)
        
        print(f"\n🎉 知识库修复完成！")
        print(f"💡 请重启后端服务以使用新的知识库")
        print(f"📊 建议测试查询: '软件工程的主要内容是什么？'")
        
    except Exception as e:
        print(f"❌ 修复过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
