# backend/process_knowledge/process_docx.py (优化版)
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import re

# --- 配置 ---
# 使用绝对路径，更加健壮
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, '..', 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, '..', 'chroma_db'))
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

def main():
    """
    处理DOCX文件，生成向量嵌入并存入ChromaDB。
    """
    print("🚀 开始处理 DOCX 文件...")

    if not os.path.exists(DOCX_SOURCE_PATH):
        print(f"❌ 错误：找不到 DOCX 文件，请确保 '{DOCX_SOURCE_PATH}' 文件存在。")
        return

    # 1. 从 DOCX 加载原始段落
    print(f"📄 正在从 DOCX 文件加载内容: {DOCX_SOURCE_PATH}")
    document = docx.Document(DOCX_SOURCE_PATH)
    raw_paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    print(f"✅ DOCX 加载完成，原始段落数: {len(raw_paragraphs)}。")

    # 2. 【优化】合并短段落，创建更有意义的文本块
    print("🧹 正在合并短段落以优化上下文...")
    MIN_CHUNK_LENGTH = 50  # 最小文本块长度，小于此长度的段落会尝试与前后合并
    merged_texts = []
    current_chunk = ""

    for para in raw_paragraphs:
        # 如果当前块加上新段落不超过 chunk_size，就合并
        if len(current_chunk) + len(para) < 800:
            current_chunk += "\n" + para
        else:
            # 否则，保存当前块，并用新段落开始一个新块
            merged_texts.append(current_chunk.strip())
            current_chunk = para
    
    # 不要忘记最后一个块
    if current_chunk:
        merged_texts.append(current_chunk.strip())

    # 将合并后的文本转换为 Document 对象
    documents = [Document(page_content=text, metadata={"source": DOCX_SOURCE_PATH}) for text in merged_texts]
    print(f"✅ 上下文优化完成，生成 {len(documents)} 个初始文本块。")

    # 3. 对可能过长的文本块进行最后的切分
    print("🔪 正在对过长的文本块进行最终切分...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "。", "！", "？", "，", "、", " "]
    )
    texts = text_splitter.split_documents(documents)
    
    if not texts:
        print("❌ 错误：文本分块后内容为空。")
        return
        
    print(f"✅ 文本分块完成，最终生成 {len(texts)} 个向量化的文本块。")

    # 4. 加载嵌入模型
    print(f"🧠 正在加载嵌入模型: {EMBEDDING_MODEL_NAME}")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    print("✅ 嵌入模型加载成功。")

    # 5. 创建并持久化向量数据库
    print(f"💾 正在创建向量数据库并存入 '{CHROMA_DB_PATH}'...")
    vectordb = Chroma.from_documents(
        documents=texts, 
        embedding=embeddings,
        persist_directory=CHROMA_DB_PATH
    )
    vectordb.persist()
    print("🎉 所有步骤完成！您的知识库已成功更新。")


if __name__ == "__main__":
    main()