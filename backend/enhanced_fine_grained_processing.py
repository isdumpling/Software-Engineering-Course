#!/usr/bin/env python3
"""
更精细化的知识库处理 - 专门捕获重要概念
"""
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import shutil
import re

# 尝试使用新版本
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain_community.vectorstores import Chroma

# 配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'chroma_db'))
EMBEDDING_MODEL_NAME = "BAAI/bge-large-zh-v1.5"

def extract_all_content_with_context(docx_path):
    """提取所有内容并保持丰富的上下文"""
    print(f"📄 正在精细化提取文档内容...")
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"原始段落数: {len(paragraphs)}")
    
    # 1. 首先寻找重要的专业术语和概念
    important_terms = [
        "两次飞跃", "质的飞跃", "认识飞跃",
        "瀑布模型", "螺旋模型", "增量模型", "原型模型",
        "软件危机", "软件工程", "软件生存周期", "软件生命周期",
        "需求分析", "系统设计", "详细设计", "编码实现", "软件测试",
        "面向对象", "OOA", "OOD", "UML", "用例图", "类图",
        "黑盒测试", "白盒测试", "单元测试", "集成测试", "系统测试",
        "软件维护", "软件质量", "软件可靠性", "软件复用"
    ]
    
    # 2. 为每个重要术语创建增强的上下文块
    enhanced_documents = []
    
    for term in important_terms:
        term_contexts = find_term_with_context(paragraphs, term)
        if term_contexts:
            for context in term_contexts:
                enhanced_doc = Document(
                    page_content=f"[关键概念-{term}] {context}",
                    metadata={
                        "source": DOCX_SOURCE_PATH,
                        "term": term,
                        "type": "concept_enhanced",
                        "priority": 1
                    }
                )
                enhanced_documents.append(enhanced_doc)
    
    print(f"✅ 为重要术语创建了 {len(enhanced_documents)} 个增强文档")
    
    # 3. 创建完整的段落级文档（保持更多上下文）
    paragraph_documents = []
    for i, para in enumerate(paragraphs):
        if len(para) > 30:  # 只保留有意义的段落
            # 为每个段落添加前后段落的上下文
            context_before = paragraphs[max(0, i-1)] if i > 0 else ""
            context_after = paragraphs[min(len(paragraphs)-1, i+1)] if i < len(paragraphs)-1 else ""
            
            enriched_content = f"{context_before}\n\n{para}\n\n{context_after}".strip()
            
            doc = Document(
                page_content=enriched_content,
                metadata={
                    "source": DOCX_SOURCE_PATH,
                    "paragraph_index": i,
                    "type": "paragraph_with_context",
                    "priority": 2
                }
            )
            paragraph_documents.append(doc)
    
    print(f"✅ 创建了 {len(paragraph_documents)} 个段落级文档")
    
    # 4. 创建更小粒度的文档块（捕获细节）
    fine_grained_documents = []
    all_text = "\n".join(paragraphs)
    
    # 使用更小的chunk_size来确保重要概念不被截断
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,  # 更小的块
        chunk_overlap=100,  # 更大的重叠
        separators=["\n\n", "。", "！", "？", "\n", "，", "；"],
        length_function=len,
    )
    
    temp_doc = Document(page_content=all_text, metadata={"source": DOCX_SOURCE_PATH})
    chunks = text_splitter.split_documents([temp_doc])
    
    for i, chunk in enumerate(chunks):
        chunk.metadata.update({
            "type": "fine_grained",
            "chunk_index": i,
            "priority": 3
        })
        fine_grained_documents.append(chunk)
    
    print(f"✅ 创建了 {len(fine_grained_documents)} 个细粒度文档")
    
    # 5. 合并所有类型的文档
    all_documents = enhanced_documents + paragraph_documents + fine_grained_documents
    print(f"📊 总计创建了 {len(all_documents)} 个文档")
    
    return all_documents

def find_term_with_context(paragraphs, term):
    """寻找包含特定术语的段落，并提供丰富的上下文"""
    contexts = []
    
    for i, para in enumerate(paragraphs):
        if term in para:
            # 获取前后各2个段落作为上下文
            start_idx = max(0, i-2)
            end_idx = min(len(paragraphs), i+3)
            
            context_paragraphs = paragraphs[start_idx:end_idx]
            full_context = "\n".join(context_paragraphs)
            
            # 确保上下文不会太长
            if len(full_context) <= 1000:
                contexts.append(full_context)
            else:
                # 如果太长，只取当前段落和前后各1个段落
                start_idx = max(0, i-1)
                end_idx = min(len(paragraphs), i+2)
                shorter_context = "\n".join(paragraphs[start_idx:end_idx])
                contexts.append(shorter_context)
    
    return contexts

def rebuild_with_fine_grained_approach():
    """使用精细化方法重建向量数据库"""
    print("🚀 开始精细化重建向量数据库...")
    
    try:
        # 1. 删除旧数据库
        if os.path.exists(CHROMA_DB_PATH):
            print(f"🗑️ 删除旧向量数据库...")
            shutil.rmtree(CHROMA_DB_PATH)
        
        # 2. 精细化提取内容
        documents = extract_all_content_with_context(DOCX_SOURCE_PATH)
        
        # 3. 加载BGE模型
        print(f"🧠 加载BGE嵌入模型...")
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
        
        # 4. 创建新的向量数据库
        print(f"💾 创建精细化向量数据库...")
        vectordb = Chroma.from_documents(
            documents=documents,
            embedding=embeddings,
            persist_directory=CHROMA_DB_PATH
        )
        
        doc_count = len(vectordb._collection.get()['documents'])
        print(f"📊 精细化数据库创建成功，包含 {doc_count} 个文档")
        
        # 5. 测试关键概念的检索
        print(f"\n🧪 测试关键概念检索...")
        retriever = vectordb.as_retriever(search_kwargs={"k": 5})
        
        test_concepts = [
            "两次飞跃",
            "软件危机", 
            "瀑布模型",
            "面向对象",
            "黑盒测试",
            "软件维护"
        ]
        
        for concept in test_concepts:
            print(f"\n🔍 测试概念: '{concept}'")
            try:
                results = retriever.invoke(concept)
                found_relevant = False
                
                for i, result in enumerate(results):
                    content = result.page_content
                    if concept in content:
                        found_relevant = True
                        doc_type = result.metadata.get('type', '未知')
                        print(f"   ✅ 结果 {i+1} ({doc_type}): ...{concept}... 找到匹配!")
                        print(f"      内容: {content[:100]}...")
                        break
                
                if not found_relevant:
                    print(f"   ⚠️ 未找到包含 '{concept}' 的直接匹配")
                    
            except Exception as e:
                print(f"   ❌ 查询失败: {e}")
        
        print(f"\n🎉 精细化重建完成！")
        return True
        
    except Exception as e:
        print(f"❌ 重建过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = rebuild_with_fine_grained_approach()
    if success:
        print(f"\n✅ 精细化重建完成！现在应该能更好地检索到各种专业概念了")
        print(f"💡 请重启后端服务以使用新的精细化数据库")
    else:
        print(f"\n❌ 重建失败，请检查错误信息")
