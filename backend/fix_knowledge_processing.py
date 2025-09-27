#!/usr/bin/env python3
"""
修复后的知识库处理脚本
解决文档分块质量问题
"""
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import re

# --- 配置 ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'chroma_db'))
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

def clean_and_filter_text(text):
    """清理和过滤文本"""
    # 移除多余的空白字符
    text = re.sub(r'\s+', ' ', text).strip()
    
    # 过滤掉太短的文本（可能只是标题或序号）
    if len(text) < 20:
        return None
        
    # 过滤掉只包含标点符号或数字的文本
    if re.match(r'^[\s\d\W]+$', text):
        return None
        
    return text

def improved_paragraph_merging(paragraphs, min_chunk_size=100, max_chunk_size=800):
    """
    改进的段落合并策略
    """
    merged_chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        cleaned_para = clean_and_filter_text(para)
        if not cleaned_para:
            continue
            
        # 如果加上新段落后不会超过最大长度，就合并
        potential_chunk = current_chunk + "\n" + cleaned_para if current_chunk else cleaned_para
        
        if len(potential_chunk) <= max_chunk_size:
            current_chunk = potential_chunk
        else:
            # 如果当前块已经足够长了，就保存它
            if len(current_chunk) >= min_chunk_size:
                merged_chunks.append(current_chunk)
            current_chunk = cleaned_para
    
    # 处理最后一个块
    if current_chunk and len(current_chunk) >= min_chunk_size:
        merged_chunks.append(current_chunk)
    
    return merged_chunks

def main():
    """
    修复后的文档处理主函数
    """
    print("🚀 开始重新处理 DOCX 文件...")

    if not os.path.exists(DOCX_SOURCE_PATH):
        print(f"❌ 错误：找不到 DOCX 文件，请确保 '{DOCX_SOURCE_PATH}' 文件存在。")
        return

    # 1. 从 DOCX 加载原始段落
    print(f"📄 正在从 DOCX 文件加载内容: {DOCX_SOURCE_PATH}")
    document = docx.Document(DOCX_SOURCE_PATH)
    raw_paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    print(f"✅ DOCX 加载完成，原始段落数: {len(raw_paragraphs)}。")

    # 显示一些段落样本
    print(f"\n📋 段落内容示例:")
    for i, para in enumerate(raw_paragraphs[:5]):
        print(f"段落 {i+1} (长度: {len(para)}): {para[:100]}...")

    # 2. 改进的段落合并策略
    print("\n🧹 正在使用改进策略合并段落...")
    merged_chunks = improved_paragraph_merging(raw_paragraphs, min_chunk_size=100, max_chunk_size=800)
    print(f"✅ 合并完成，生成 {len(merged_chunks)} 个文本块。")

    # 显示合并后的文本块示例
    print(f"\n📊 合并后文本块示例:")
    for i, chunk in enumerate(merged_chunks[:3]):
        print(f"\n--- 文本块 {i+1} (长度: {len(chunk)}) ---")
        print(f"{chunk[:200]}...")

    # 3. 转换为 Document 对象
    documents = [Document(page_content=chunk, metadata={"source": DOCX_SOURCE_PATH, "chunk_id": i}) 
                for i, chunk in enumerate(merged_chunks)]
    
    # 4. 进一步文本分割（如果需要）
    print("\n🔪 正在进行最终文本分割...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "。", "！", "？", "\n", "，", "、", " "],
        length_function=len,
    )
    final_texts = text_splitter.split_documents(documents)
    
    if not final_texts:
        print("❌ 错误：文本分块后内容为空。")
        return
        
    print(f"✅ 最终分割完成，共生成 {len(final_texts)} 个文档块。")

    # 显示最终文档示例
    print(f"\n📝 最终文档示例:")
    for i, doc in enumerate(final_texts[:3]):
        print(f"\n--- 最终文档 {i+1} (长度: {len(doc.page_content)}) ---")
        print(f"内容: {doc.page_content[:150]}...")
        print(f"元数据: {doc.metadata}")

    # 5. 加载嵌入模型
    print(f"\n🧠 正在加载嵌入模型: {EMBEDDING_MODEL_NAME}")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    print("✅ 嵌入模型加载成功。")

    # 6. 删除旧的向量数据库
    import shutil
    if os.path.exists(CHROMA_DB_PATH):
        print(f"🗑️ 正在删除旧的向量数据库...")
        shutil.rmtree(CHROMA_DB_PATH)

    # 7. 创建并持久化新的向量数据库
    print(f"💾 正在创建新的向量数据库并存入 '{CHROMA_DB_PATH}'...")
    vectordb = Chroma.from_documents(
        documents=final_texts, 
        embedding=embeddings,
        persist_directory=CHROMA_DB_PATH
    )
    vectordb.persist()
    print("🎉 知识库重新构建完成！")

    # 8. 测试检索功能
    print(f"\n🔍 测试新知识库的检索功能...")
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})
    test_queries = ["软件开发生命周期", "测试方法", "软件工程"]
    
    for query in test_queries:
        print(f"\n查询: '{query}'")
        results = retriever.invoke(query)  # 使用新的invoke方法
        for i, result in enumerate(results):
            print(f"  结果 {i+1}: {result.page_content[:100]}...")

if __name__ == "__main__":
    main()
