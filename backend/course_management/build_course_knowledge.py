#!/usr/bin/env python3
"""
单课程知识库构建脚本
"""
import os
import sys
import docx
import shutil
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# 添加上级目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain_community.vectorstores import Chroma
from course_config import CourseManager, COURSE_CONFIG

# 全局配置
EMBEDDING_MODEL_NAME = "BAAI/bge-large-zh-v1.5"

def extract_course_content(course_id, docx_path):
    """提取指定课程的内容"""
    print(f"📄 正在提取课程内容: {course_id}")
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"课程教材文件不存在: {docx_path}")
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"原始段落数: {len(paragraphs)}")
    
    # 获取课程配置
    course_config = COURSE_CONFIG.get(course_id, {})
    important_terms = course_config.get("important_terms", [])
    course_name = course_config.get("name", course_id)
    
    print(f"课程: {course_name}")
    print(f"重要术语数量: {len(important_terms)}")
    
    # 1. 为重要术语创建增强文档
    enhanced_documents = []
    for term in important_terms:
        term_contexts = find_term_with_context(paragraphs, term)
        if term_contexts:
            for context in term_contexts:
                enhanced_doc = Document(
                    page_content=f"[{course_name}-关键概念-{term}] {context}",
                    metadata={
                        "source": docx_path,
                        "course_id": course_id,
                        "course_name": course_name,
                        "term": term,
                        "type": "concept_enhanced",
                        "priority": 1
                    }
                )
                enhanced_documents.append(enhanced_doc)
    
    print(f"✅ 为重要术语创建了 {len(enhanced_documents)} 个增强文档")
    
    # 2. 创建段落级文档（保持上下文）
    paragraph_documents = []
    for i, para in enumerate(paragraphs):
        if len(para) > 30:
            # 添加前后段落的上下文
            context_before = paragraphs[max(0, i-1)] if i > 0 else ""
            context_after = paragraphs[min(len(paragraphs)-1, i+1)] if i < len(paragraphs)-1 else ""
            
            enriched_content = f"{context_before}\n\n{para}\n\n{context_after}".strip()
            
            doc = Document(
                page_content=f"[{course_name}-段落] {enriched_content}",
                metadata={
                    "source": docx_path,
                    "course_id": course_id,
                    "course_name": course_name,
                    "paragraph_index": i,
                    "type": "paragraph_with_context",
                    "priority": 2
                }
            )
            paragraph_documents.append(doc)
    
    print(f"✅ 创建了 {len(paragraph_documents)} 个段落级文档")
    
    # 3. 创建细粒度文档
    fine_grained_documents = []
    all_text = "\n".join(paragraphs)
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=100,
        separators=["\n\n", "。", "！", "？", "\n", "，", "；"],
        length_function=len,
    )
    
    temp_doc = Document(page_content=all_text, metadata={"source": docx_path})
    chunks = text_splitter.split_documents([temp_doc])
    
    for i, chunk in enumerate(chunks):
        chunk.page_content = f"[{course_name}-细节] {chunk.page_content}"
        chunk.metadata.update({
            "course_id": course_id,
            "course_name": course_name,
            "type": "fine_grained",
            "chunk_index": i,
            "priority": 3
        })
        fine_grained_documents.append(chunk)
    
    print(f"✅ 创建了 {len(fine_grained_documents)} 个细粒度文档")
    
    # 合并所有文档
    all_documents = enhanced_documents + paragraph_documents + fine_grained_documents
    print(f"📊 总计创建了 {len(all_documents)} 个文档")
    
    return all_documents

def find_term_with_context(paragraphs, term):
    """寻找包含特定术语的段落，并提供上下文"""
    contexts = []
    
    for i, para in enumerate(paragraphs):
        if term in para:
            # 获取前后各2个段落作为上下文
            start_idx = max(0, i-2)
            end_idx = min(len(paragraphs), i+3)
            
            context_paragraphs = paragraphs[start_idx:end_idx]
            full_context = "\n".join(context_paragraphs)
            
            # 确保上下文不会太长
            if len(full_context) <= 1000:
                contexts.append(full_context)
            else:
                # 如果太长，只取当前段落和前后各1个段落
                start_idx = max(0, i-1)
                end_idx = min(len(paragraphs), i+2)
                shorter_context = "\n".join(paragraphs[start_idx:end_idx])
                contexts.append(shorter_context)
    
    return contexts

def build_course_vectordb(course_id):
    """构建指定课程的向量数据库"""
    print(f"🚀 开始构建课程向量数据库: {course_id}")
    
    try:
        manager = CourseManager()
        
        # 检查教材文件
        doc_path = manager.get_course_doc_path(course_id)
        if not doc_path or not os.path.exists(doc_path):
            print(f"❌ 课程教材文件不存在: {doc_path}")
            print(f"请确保文件存在: {doc_path}")
            return False
        
        # 获取向量数据库路径
        db_path = manager.get_course_vector_db_path(course_id)
        
        # 删除旧的向量数据库
        if os.path.exists(db_path):
            print(f"🗑️ 删除旧的向量数据库: {db_path}")
            shutil.rmtree(db_path)
        
        # 确保父目录存在
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        
        # 提取文档内容
        documents = extract_course_content(course_id, doc_path)
        
        if not documents:
            print(f"❌ 未能提取到有效文档内容")
            return False
        
        # 加载嵌入模型
        print(f"🧠 加载BGE嵌入模型: {EMBEDDING_MODEL_NAME}")
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
        
        # 创建向量数据库
        print(f"💾 创建向量数据库: {db_path}")
        vectordb = Chroma.from_documents(
            documents=documents,
            embedding=embeddings,
            persist_directory=db_path
        )
        
        # 验证结果
        doc_count = len(vectordb._collection.get()['documents'])
        print(f"📊 向量数据库创建成功，包含 {doc_count} 个文档")
        
        # 测试检索
        print(f"🧪 测试课程概念检索...")
        course_config = COURSE_CONFIG.get(course_id, {})
        test_terms = course_config.get("important_terms", [])[:3]  # 测试前3个重要术语
        
        retriever = vectordb.as_retriever(search_kwargs={"k": 3})
        
        for term in test_terms:
            print(f"\n🔍 测试术语: '{term}'")
            try:
                results = retriever.invoke(term)
                found_relevant = False
                
                for i, result in enumerate(results):
                    if term in result.page_content:
                        found_relevant = True
                        doc_type = result.metadata.get('type', '未知')
                        print(f"   ✅ 结果 {i+1} ({doc_type}): 找到匹配!")
                        break
                
                if not found_relevant:
                    print(f"   ⚠️ 未找到包含 '{term}' 的直接匹配")
                    
            except Exception as e:
                print(f"   ❌ 查询失败: {e}")
        
        print(f"\n🎉 课程 {course_id} 的向量数据库构建完成！")
        return True
        
    except Exception as e:
        print(f"❌ 构建过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    if len(sys.argv) != 2:
        print("用法: python build_course_knowledge.py <course_id>")
        print("\n支持的课程ID:")
        manager = CourseManager()
        for course_id in manager.get_all_courses():
            course_name = COURSE_CONFIG[course_id]["name"]
            print(f"  - {course_id} ({course_name})")
        sys.exit(1)
    
    course_id = sys.argv[1]
    
    # 验证课程ID
    if course_id not in COURSE_CONFIG:
        print(f"❌ 不支持的课程ID: {course_id}")
        print("\n支持的课程ID:")
        for cid in COURSE_CONFIG.keys():
            course_name = COURSE_CONFIG[cid]["name"]
            print(f"  - {cid} ({course_name})")
        sys.exit(1)
    
    # 构建向量数据库
    success = build_course_vectordb(course_id)
    
    if success:
        print(f"\n✅ 课程 {course_id} 的知识库构建成功！")
        print(f"💡 请重启后端服务以使用新的知识库")
    else:
        print(f"\n❌ 课程 {course_id} 的知识库构建失败")
        sys.exit(1)

if __name__ == "__main__":
    main()
