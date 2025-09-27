#!/usr/bin/env python3
"""
针对性修复检索问题 - 确保生命周期内容被正确检索
"""
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
import re
import shutil

# 尝试使用新版本
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

# 配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'chroma_db'))
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

def extract_lifecycle_focused_content(docx_path):
    """重点提取软件生命周期相关内容"""
    print(f"📄 正在提取生命周期重点内容...")
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"总段落数: {len(paragraphs)}")
    
    # 定义重要概念的关键词
    important_concepts = {
        'software_lifecycle': ['生命周期', '生存周期', '开发周期', '软件过程'],
        'development_phases': ['需求分析', '系统设计', '编码', '测试', '维护', '运行'],
        'development_models': ['瀑布模型', '螺旋模型', '原型模型', '增量模型', '构件'],
        'software_engineering': ['软件工程', '软件开发', '开发过程', '开发方法'],
        'analysis_design': ['面向对象', 'OOA', 'OOD', 'UML', '模块化'],
        'testing_quality': ['软件测试', '质量保证', '可靠性', '纠错'],
        'project_management': ['项目管理', '软件管理', '开发管理']
    }
    
    # 按重要性分类段落
    categorized_content = {category: [] for category in important_concepts.keys()}
    uncategorized_content = []
    
    for i, para in enumerate(paragraphs):
        if len(para) < 20:  # 跳过太短的段落
            continue
            
        para_category = None
        max_matches = 0
        
        # 找出最匹配的类别
        for category, keywords in important_concepts.items():
            matches = sum(1 for kw in keywords if kw in para)
            if matches > max_matches:
                max_matches = matches
                para_category = category
        
        # 分类存储
        if para_category and max_matches > 0:
            categorized_content[para_category].append({
                'index': i,
                'content': para,
                'matches': max_matches
            })
        else:
            uncategorized_content.append({
                'index': i,
                'content': para,
                'matches': 0
            })
    
    # 显示分类结果
    print(f"\n📊 内容分类结果:")
    for category, items in categorized_content.items():
        if items:
            print(f"  {category}: {len(items)} 段落")
            # 显示最重要的内容
            items.sort(key=lambda x: x['matches'], reverse=True)
            for item in items[:2]:  # 显示前2个最匹配的
                print(f"    - 段落{item['index']} (匹配{item['matches']}个关键词): {item['content'][:80]}...")
    
    print(f"  未分类: {len(uncategorized_content)} 段落")
    
    return categorized_content, uncategorized_content

def create_priority_documents(categorized_content, uncategorized_content):
    """创建有优先级的文档"""
    print(f"\n📚 正在创建优先级文档...")
    
    documents = []
    
    # 定义优先级（数字越小优先级越高）
    category_priority = {
        'software_lifecycle': 1,
        'development_phases': 2,
        'development_models': 3,
        'software_engineering': 4,
        'analysis_design': 5,
        'testing_quality': 6,
        'project_management': 7
    }
    
    # 按优先级处理分类内容
    for category, priority in sorted(category_priority.items(), key=lambda x: x[1]):
        items = categorized_content.get(category, [])
        if not items:
            continue
            
        print(f"  处理 {category} (优先级 {priority}): {len(items)} 个项目")
        
        # 按匹配度排序
        items.sort(key=lambda x: x['matches'], reverse=True)
        
        # 对每个类别的内容进行合理分组
        for item in items:
            content = item['content']
            
            # 为重要内容添加标签
            if priority <= 3:  # 高优先级内容
                tagged_content = f"[重要-{category}] {content}"
            else:
                tagged_content = f"[{category}] {content}"
            
            # 如果内容太长，进行分割
            if len(tagged_content) > 800:
                # 使用文本分割器
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=700,
                    chunk_overlap=100,
                    separators=["\n\n", "。", "！", "？", "\n", "，"],
                    length_function=len,
                )
                
                temp_doc = Document(page_content=tagged_content, metadata={
                    "source": DOCX_SOURCE_PATH,
                    "category": category,
                    "priority": priority,
                    "original_index": item['index'],
                    "keyword_matches": item['matches']
                })
                
                sub_docs = text_splitter.split_documents([temp_doc])
                documents.extend(sub_docs)
            else:
                # 直接创建文档
                doc = Document(
                    page_content=tagged_content,
                    metadata={
                        "source": DOCX_SOURCE_PATH,
                        "category": category,
                        "priority": priority,
                        "original_index": item['index'],
                        "keyword_matches": item['matches'],
                        "length": len(tagged_content)
                    }
                )
                documents.append(doc)
    
    # 处理未分类内容（低优先级）
    print(f"  处理未分类内容: {len(uncategorized_content)} 个项目")
    for item in uncategorized_content:
        if len(item['content']) > 50:  # 只保留足够长的内容
            doc = Document(
                page_content=f"[其他] {item['content']}",
                metadata={
                    "source": DOCX_SOURCE_PATH,
                    "category": "other",
                    "priority": 8,
                    "original_index": item['index'],
                    "keyword_matches": 0,
                    "length": len(item['content'])
                }
            )
            documents.append(doc)
    
    print(f"✅ 创建了 {len(documents)} 个优先级文档")
    
    # 显示一些示例
    print(f"\n📋 高优先级文档示例:")
    high_priority_docs = [d for d in documents if d.metadata.get('priority', 8) <= 3]
    for i, doc in enumerate(high_priority_docs[:3]):
        print(f"  文档 {i+1} (优先级 {doc.metadata['priority']}):")
        print(f"    类别: {doc.metadata['category']}")
        print(f"    内容: {doc.page_content[:100]}...")
    
    return documents

def rebuild_optimized_vectordb(documents):
    """重建优化的向量数据库"""
    print(f"\n🚀 正在重建优化的向量数据库...")
    
    # 删除旧数据库
    if os.path.exists(CHROMA_DB_PATH):
        print(f"🗑️ 删除旧向量数据库...")
        shutil.rmtree(CHROMA_DB_PATH)
    
    # 加载嵌入模型
    print(f"🧠 加载嵌入模型: {EMBEDDING_MODEL_NAME}")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    
    # 创建新的向量数据库
    print(f"💾 创建新向量数据库...")
    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=CHROMA_DB_PATH
    )
    
    print(f"✅ 优化向量数据库重建完成！")
    return vectordb

def test_lifecycle_queries(vectordb):
    """专门测试生命周期相关查询"""
    print(f"\n🧪 专门测试生命周期查询...")
    
    lifecycle_queries = [
        ("软件开发生命周期有哪些阶段", "应该返回生命周期各个阶段的详细信息"),
        ("软件生存周期", "应该返回生存周期的定义和阶段划分"),
        ("瀑布模型", "应该返回瀑布模型的特点和流程"),
        ("需求分析阶段的任务", "应该返回需求分析的具体任务"),
        ("软件开发过程", "应该返回软件开发的整体过程"),
        ("软件工程的基本概念", "应该返回软件工程的定义和原理")
    ]
    
    # 测试不同的检索参数
    test_configs = [
        {"k": 3, "description": "标准检索(top-3)"},
        {"k": 5, "description": "扩展检索(top-5)"},
    ]
    
    for config in test_configs:
        print(f"\n" + "="*60)
        print(f"🔧 测试配置: {config['description']}")
        
        # 正确的参数传递方式
        search_kwargs = {"k": config["k"]}
        retriever = vectordb.as_retriever(search_kwargs=search_kwargs)
        
        for query, expected in lifecycle_queries:
            print(f"\n🔍 查询: '{query}'")
            print(f"📝 期望: {expected}")
            
            try:
                results = retriever.invoke(query)
                print(f"✅ 检索到 {len(results)} 个结果")
                
                for i, result in enumerate(results):
                    print(f"\n📄 结果 {i+1}:")
                    print(f"   类别: {result.metadata.get('category', '未知')}")
                    print(f"   优先级: {result.metadata.get('priority', '未知')}")
                    print(f"   关键词匹配: {result.metadata.get('keyword_matches', 0)}")
                    print(f"   内容: {result.page_content[:150]}...")
                    
                    # 计算查询相关性
                    query_words = set(query.lower().split())
                    content_words = set(result.page_content.lower().split())
                    overlap = len(query_words.intersection(content_words))
                    relevance = overlap / len(query_words) if query_words else 0
                    print(f"   相关性得分: {relevance:.2f}")
                
                # 分析结果质量
                high_priority_count = sum(1 for r in results if r.metadata.get('priority', 8) <= 3)
                lifecycle_count = sum(1 for r in results if 'lifecycle' in r.metadata.get('category', ''))
                
                print(f"\n📊 结果质量分析:")
                print(f"   高优先级结果: {high_priority_count}/{len(results)}")
                print(f"   生命周期相关: {lifecycle_count}/{len(results)}")
                
            except Exception as e:
                print(f"❌ 查询失败: {e}")

def main():
    """主函数"""
    print("🚀 开始针对性修复检索问题...")
    
    try:
        # 1. 提取生命周期重点内容
        categorized_content, uncategorized_content = extract_lifecycle_focused_content(DOCX_SOURCE_PATH)
        
        # 2. 创建优先级文档
        documents = create_priority_documents(categorized_content, uncategorized_content)
        
        if not documents:
            print("❌ 未能创建有效文档")
            return
        
        # 3. 重建优化向量数据库
        vectordb = rebuild_optimized_vectordb(documents)
        
        # 4. 专门测试生命周期查询
        test_lifecycle_queries(vectordb)
        
        print(f"\n🎉 针对性修复完成！")
        print(f"💡 请重启后端服务，然后测试查询: '软件开发生命周期的主要阶段'")
        
    except Exception as e:
        print(f"❌ 修复过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
