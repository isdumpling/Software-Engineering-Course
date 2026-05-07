"""课程知识库管理服务 —— 由原 course_management/build_course_knowledge.py 重构"""

import os
import shutil
from typing import Callable, Dict, List, Optional

import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

from course_management.course_config import CourseManager, COURSE_CONFIG

EMBEDDING_MODEL_NAME = "BAAI/bge-large-zh-v1.5"


def _log(log_cb: Optional[Callable[[str], None]], message: str):
    print(message)
    if log_cb:
        log_cb(message)


def find_term_with_context(paragraphs: List[str], term: str) -> List[str]:
    contexts = []
    for i, para in enumerate(paragraphs):
        if term in para:
            start_idx = max(0, i - 2)
            end_idx = min(len(paragraphs), i + 3)
            full_context = "\n".join(paragraphs[start_idx:end_idx])
            if len(full_context) <= 1000:
                contexts.append(full_context)
            else:
                contexts.append("\n".join(paragraphs[max(0, i - 1):min(len(paragraphs), i + 2)]))
    return contexts


def extract_course_content(
    course_id: str,
    docx_path: str,
    log_cb: Optional[Callable[[str], None]] = None
) -> List[Document]:
    _log(log_cb, f"正在提取课程内容: {course_id}")

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"课程教材文件不存在: {docx_path}")

    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    _log(log_cb, f"原始段落数: {len(paragraphs)}")

    course_config = COURSE_CONFIG.get(course_id, {})
    important_terms = course_config.get("important_terms", [])
    course_name = course_config.get("name", course_id)

    _log(log_cb, f"课程: {course_name}  重要术语: {len(important_terms)} 个")

    # 概念增强文档
    enhanced_documents = []
    for term in important_terms:
        for context in find_term_with_context(paragraphs, term):
            enhanced_documents.append(Document(
                page_content=f"[{course_name}-关键概念-{term}] {context}",
                metadata={
                    "source": docx_path, "course_id": course_id,
                    "course_name": course_name, "term": term,
                    "type": "concept_enhanced", "priority": 1,
                },
            ))
    _log(log_cb, f"概念增强文档: {len(enhanced_documents)} 个")

    # 段落级文档（带上下文窗口）
    paragraph_documents = []
    for i, para in enumerate(paragraphs):
        if len(para) > 30:
            before = paragraphs[max(0, i - 1)] if i > 0 else ""
            after = paragraphs[min(len(paragraphs) - 1, i + 1)] if i < len(paragraphs) - 1 else ""
            enriched = f"{before}\n\n{para}\n\n{after}".strip()
            paragraph_documents.append(Document(
                page_content=f"[{course_name}-段落] {enriched}",
                metadata={
                    "source": docx_path, "course_id": course_id,
                    "course_name": course_name, "paragraph_index": i,
                    "type": "paragraph_with_context", "priority": 2,
                },
            ))
    _log(log_cb, f"段落级文档: {len(paragraph_documents)} 个")

    # 细粒度切分
    all_text = "\n".join(paragraphs)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300, chunk_overlap=100,
        separators=["\n\n", "。", "！", "？", "\n", "，", "；"],
        length_function=len,
    )
    temp_doc = Document(page_content=all_text, metadata={"source": docx_path})
    fine_docs = []
    for i, chunk in enumerate(splitter.split_documents([temp_doc])):
        chunk.page_content = f"[{course_name}-细节] {chunk.page_content}"
        chunk.metadata.update({
            "course_id": course_id, "course_name": course_name,
            "type": "fine_grained", "chunk_index": i, "priority": 3,
        })
        fine_docs.append(chunk)
    _log(log_cb, f"细粒度文档: {len(fine_docs)} 个")

    all_docs = enhanced_documents + paragraph_documents + fine_docs
    _log(log_cb, f"总计: {len(all_docs)} 个文档块")
    return all_docs


def build_course_vectordb(
    course_id: str,
    log_cb: Optional[Callable[[str], None]] = None
) -> Dict[str, object]:
    """构建指定课程的向量数据库。返回 {"success": bool, "error"/data...}"""
    try:
        mgr = CourseManager()
        if course_id not in COURSE_CONFIG:
            return {"success": False, "error": f"不支持的课程 ID: {course_id}"}

        cfg = COURSE_CONFIG[course_id]
        doc_path = mgr.get_course_doc_path(course_id)
        _log(log_cb, f"开始构建: {course_id} - {cfg['name']}")
        _log(log_cb, f"教材路径: {doc_path}")

        if not doc_path or not os.path.exists(doc_path):
            return {"success": False, "error": f"教材文件不存在: {doc_path}"}

        db_path = mgr.get_course_vector_db_path(course_id)
        if os.path.exists(db_path):
            _log(log_cb, f"删除旧向量库: {db_path}")
            shutil.rmtree(db_path)
        os.makedirs(os.path.dirname(db_path), exist_ok=True)

        documents = extract_course_content(course_id, doc_path, log_cb=log_cb)
        if not documents:
            return {"success": False, "error": "未能提取到有效文档内容"}

        _log(log_cb, f"加载嵌入模型: {EMBEDDING_MODEL_NAME}")
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)

        _log(log_cb, f"创建向量库: {db_path}")
        vectordb = Chroma.from_documents(
            documents=documents, embedding=embeddings, persist_directory=db_path,
        )
        doc_count = len(vectordb._collection.get()["documents"])
        _log(log_cb, f"向量库创建完成: {doc_count} 个文档块")

        return {
            "success": True, "course_id": course_id,
            "course_name": cfg["name"], "document_count": doc_count,
            "vector_db_path": db_path, "source_file": doc_path,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_all_course_status(ai_instance=None) -> List[Dict]:
    mgr = CourseManager()
    stats = mgr.list_all_course_stats()
    result = []
    for course_id, item in stats.items():
        cfg = COURSE_CONFIG.get(course_id, {})
        loaded = False
        if ai_instance is not None:
            loaded = course_id in getattr(ai_instance, "course_retrievers", {})
        result.append({
            "course_id": course_id,
            "course_name": cfg.get("name", item.get("course_name", course_id)),
            "doc_file": cfg.get("doc_file", ""),
            "doc_exists": item.get("doc_exists", False),
            "doc_size": item.get("doc_size", 0),
            "vector_db_exists": item.get("vector_db_exists", False),
            "vector_db_size": item.get("vector_db_size", 0),
            "document_count": item.get("document_count", 0),
            "loaded": loaded,
        })
    return result


def save_course_material(course_id: str, file_bytes: bytes, filename: str) -> Dict:
    mgr = CourseManager()
    if course_id not in COURSE_CONFIG:
        return {"success": False, "error": f"不支持的课程 ID: {course_id}"}
    if not filename.lower().endswith(".docx"):
        return {"success": False, "error": "仅支持上传 .docx 教材文件"}

    doc_path = mgr.get_course_doc_path(course_id)
    if not doc_path:
        return {"success": False, "error": "无法确定课程教材保存路径"}

    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    with open(doc_path, "wb") as f:
        f.write(file_bytes)

    return {
        "success": True, "course_id": course_id,
        "filename": filename, "saved_path": doc_path, "size": len(file_bytes),
    }
