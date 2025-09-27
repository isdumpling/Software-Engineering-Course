#!/usr/bin/env python3
"""
增强版知识库处理 - 改进检索质量
"""
import os
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import re

# 配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_SOURCE_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'knowledge_base', 'se.docx'))
CHROMA_DB_PATH = os.path.normpath(os.path.join(SCRIPT_DIR, 'chroma_db'))
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

def extract_and_analyze_content(docx_path):
    """提取并分析DOCX内容"""
    print(f"📄 正在提取和分析DOCX内容...")
    
    doc = docx.Document(docx_path)
    
    # 提取所有文本内容，包括标题信息
    all_content = []
    current_section = ""
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # 检查是否是标题（通常较短且可能包含特定格式）
        if len(text) < 100 and any(keyword in text for keyword in ['章', '节', '第', '1.', '2.', '3.', '一、', '二、']):
            current_section = text
            print(f"📚 发现章节: {current_section}")
        else:
            # 将内容与章节信息结合
            if current_section and len(text) > 50:
                enhanced_text = f"[{current_section}] {text}"
                all_content.append(enhanced_text)
            elif len(text) > 50:
                all_content.append(text)
    
    print(f"✅ 提取完成，共获得 {len(all_content)} 段有效内容")
    return all_content

def create_smart_chunks(content_list, target_size=600, overlap=100):
    """智能创建文档块"""
    print(f"🧠 正在智能创建文档块...")
    
    chunks = []
    current_chunk = ""
    
    for content in content_list:
        # 如果加上新内容不会超出目标大小，就合并
        if len(current_chunk) + len(content) + 1 <= target_size:
            if current_chunk:
                current_chunk += "\n" + content
            else:
                current_chunk = content
        else:
            # 保存当前块（如果足够大）
            if len(current_chunk) > 100:
                chunks.append(current_chunk.strip())
            
            # 开始新块
            current_chunk = content
    
    # 处理最后一个块
    if current_chunk and len(current_chunk) > 100:
        chunks.append(current_chunk.strip())
    
    print(f"✅ 创建了 {len(chunks)} 个智能文档块")
    
    # 显示块的内容摘要
    for i, chunk in enumerate(chunks[:5]):
        print(f"\n📄 块 {i+1} (长度: {len(chunk)}):")
        # 提取关键信息
        lines = chunk.split('\n')
        if lines:
            first_line = lines[0][:100] + "..." if len(lines[0]) > 100 else lines[0]
            print(f"   开始: {first_line}")
        
        # 检查关键词
        keywords = ['软件开发', '生命周期', '测试', '需求', '设计', '编码', '维护', '管理']
        found = [kw for kw in keywords if kw in chunk]
        if found:
            print(f"   关键词: {found}")
    
    return chunks

def add_keyword_metadata(chunks):
    """为文档块添加关键词元数据"""
    print(f"🏷️ 正在为文档块添加关键词元数据...")
    
    # 定义关键词分类
    keyword_categories = {
        '软件开发生命周期': ['生命周期', '开发过程', '软件过程', '开发阶段', '项目管理'],
        '需求分析': ['需求', '分析', '用户需求', '功能需求', '需求获取'],
        '系统设计': ['设计', '架构', '模块', '接口', '数据结构'],
        '编码实现': ['编码', '编程', '实现', '程序', '代码'],
        '软件测试': ['测试', '验证', '确认', '调试', '质量保证'],
        '项目管理': ['管理', '计划', '进度', '成本', '风险'],
        '软件质量': ['质量', '可靠性', '性能', '安全性', '可维护性'],
        '面向对象': ['对象', '类', '继承', '封装', '多态']
    }
    
    documents = []
    for i, chunk in enumerate(chunks):
        # 识别文档块的主要主题
        topics = []
        for topic, keywords in keyword_categories.items():
            if any(kw in chunk for kw in keywords):
                topics.append(topic)
        
        metadata = {
            "source": DOCX_SOURCE_PATH,
            "chunk_id": i,
            "topics": topics,
            "length": len(chunk)
        }
        
        documents.append(Document(page_content=chunk, metadata=metadata))
        
        if topics:
            print(f"   块 {i+1}: {topics}")
    
    return documents

def create_enhanced_vectordb(documents):
    """创建增强的向量数据库"""
    print(f"🚀 正在创建增强的向量数据库...")
    
    # 删除旧数据库
    import shutil
    if os.path.exists(CHROMA_DB_PATH):
        shutil.rmtree(CHROMA_DB_PATH)
    
    # 加载嵌入模型
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    
    # 创建向量数据库
    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=CHROMA_DB_PATH
    )
    
    print(f"✅ 向量数据库创建完成！")
    return vectordb

def comprehensive_test(vectordb):
    """综合测试检索效果"""
    print(f"\n🧪 进行综合检索测试...")
    
    test_queries = [
        "软件开发生命周期有哪些阶段",
        "软件测试的方法和策略",
        "需求分析的主要任务",
        "面向对象的基本概念",
        "软件质量的评价标准",
        "项目管理的重要性"
    ]
    
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})
    
    for query in test_queries:
        print(f"\n" + "="*50)
        print(f"🔍 查询: {query}")
        
        try:
            results = retriever.invoke(query)
            
            for i, result in enumerate(results):
                print(f"\n📄 结果 {i+1}:")
                print(f"   内容: {result.page_content[:150]}...")
                print(f"   主题: {result.metadata.get('topics', '未分类')}")
                
                # 简单相关性检查
                query_words = set(query.lower().split())
                content_words = set(result.page_content.lower().split())
                overlap = len(query_words.intersection(content_words))
                print(f"   词汇重叠: {overlap} 个关键词")
                
        except Exception as e:
            print(f"❌ 查询失败: {e}")

def main():
    """主函数"""
    print("🚀 开始增强版知识库处理...")
    
    try:
        # 1. 提取和分析内容
        content_list = extract_and_analyze_content(DOCX_SOURCE_PATH)
        
        if not content_list:
            print("❌ 未能提取到有效内容")
            return
        
        # 2. 创建智能文档块
        chunks = create_smart_chunks(content_list)
        
        # 3. 添加元数据
        documents = add_keyword_metadata(chunks)
        
        # 4. 创建向量数据库
        vectordb = create_enhanced_vectordb(documents)
        
        # 5. 综合测试
        comprehensive_test(vectordb)
        
        print(f"\n🎉 增强版知识库处理完成！")
        print(f"💡 建议：重启后端服务以使用新的知识库")
        
    except Exception as e:
        print(f"❌ 处理过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
